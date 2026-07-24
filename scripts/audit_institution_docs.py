from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


ROOT = Path("/Users/prv15/Downloads/SRPB Group of Institutions - For Software")
TARGETS = [
    ROOT / "SRPB Degree College of Education/Guide Plan Academic Process -UG.docx",
    ROOT / "SRPB Degree College of Education/4. UG - University Exam Form.pdf",
    ROOT / "SRPB Degree College of Education/3. UG - CIA Form + Admit Card.pdf",
    ROOT / "SRPB Degree College of Education/UG - CIA 3rd Sem. Marksheet.xlsx",
    ROOT / "SRPB Degree College of Education/UG - Admission Register 2025-2029.xlsx",
    ROOT / "SRP Teachers' Training College/B.Ed. - Details of Students.xlsx",
    ROOT / "SRP Teachers' Training College/D.El.Ed. - Details of Students.xlsx",
]


def clean(value: object) -> str:
    return " ".join(str(value or "").split())


for path in TARGETS:
    print(f"\n### {path.name}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        document = Document(path)
        lines = [clean(paragraph.text) for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append(" | ".join(clean(cell.text) for cell in row.cells))
        print("\n".join(line for line in lines if line)[:16000])
    elif suffix == ".pdf":
        reader = PdfReader(path)
        text = "\n".join((page.extract_text() or "") for page in reader.pages[:8])
        print(text[:16000])
    elif suffix == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=False)
        for sheet in workbook.worksheets:
            print(f"-- Sheet: {sheet.title} ({sheet.max_row}x{sheet.max_column})")
            emitted = 0
            for row in sheet.iter_rows(values_only=True):
                values = [clean(value) for value in row]
                if not any(values):
                    continue
                print(" | ".join(values[:24]))
                emitted += 1
                if emitted >= 30:
                    break
